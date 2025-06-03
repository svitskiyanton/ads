#!/usr/bin/env python3
"""
Test script to debug .docx text extraction and OpenAI parameter extraction
"""

import os
import io
import tempfile
from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
import openai
import config

def setup_drive_api():
    """Setup Google Drive API authentication"""
    creds = None
    if os.path.exists('token.json'):
        creds = Credentials.from_authorized_user_file('token.json')
    
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
    
    service = build('drive', 'v3', credentials=creds)
    print("✅ Google Drive API connected successfully!")
    return service

def test_docx_extraction(service, file_id, filename):
    """Test .docx text extraction"""
    print(f"\n🧪 Testing .docx extraction for: {filename}")
    print("=" * 50)
    
    try:
        # Download the .docx file
        print("📥 Downloading .docx file...")
        request = service.files().get_media(fileId=file_id)
        
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, request)
        
        done = False
        while done is False:
            status, done = downloader.next_chunk()
            print(f"   Download progress: {int(status.progress() * 100)}%")
        
        print(f"✅ Downloaded {len(fh.getvalue())} bytes")
        
        # Save to temporary file for inspection
        temp_path = os.path.join(tempfile.gettempdir(), f"test_{filename}")
        with open(temp_path, 'wb') as f:
            f.write(fh.getvalue())
        print(f"💾 Saved to: {temp_path}")
        
        # Try to extract text using python-docx
        print("\n📄 Extracting text using python-docx...")
        fh.seek(0)
        
        try:
            from docx import Document
            doc = Document(fh)
            
            print(f"📊 Document statistics:")
            print(f"   Paragraphs: {len(doc.paragraphs)}")
            
            # Extract text from paragraphs
            text_content = []
            for i, paragraph in enumerate(doc.paragraphs):
                para_text = paragraph.text.strip()
                print(f"   Paragraph {i+1}: '{para_text[:100]}{'...' if len(para_text) > 100 else ''}' ({len(para_text)} chars)")
                if para_text:
                    text_content.append(para_text)
            
            # Try extracting from tables too
            print(f"   Tables: {len(doc.tables)}")
            for i, table in enumerate(doc.tables):
                print(f"   Table {i+1} rows: {len(table.rows)}")
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            text_content.append(cell_text)
                            print(f"     Cell: '{cell_text[:50]}{'...' if len(cell_text) > 50 else ''}' ({len(cell_text)} chars)")
            
            final_text = '\n'.join(text_content)
            print(f"\n✅ Total extracted text: {len(final_text)} characters")
            
            if final_text:
                print("\n📝 First 500 characters of extracted text:")
                print("-" * 50)
                print(final_text[:500])
                print("-" * 50)
                return final_text
            else:
                print("❌ No text extracted from .docx file")
                
                # Try alternative method - read as XML
                print("\n🔍 Attempting alternative extraction...")
                fh.seek(0)
                import zipfile
                
                with zipfile.ZipFile(fh, 'r') as zip_file:
                    print(f"   ZIP contents: {zip_file.namelist()}")
                    
                    if 'word/document.xml' in zip_file.namelist():
                        xml_content = zip_file.read('word/document.xml').decode('utf-8')
                        print(f"   XML length: {len(xml_content)} chars")
                        print(f"   XML preview: {xml_content[:200]}...")
                        
                        # Basic XML text extraction
                        import re
                        text_matches = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', xml_content)
                        if text_matches:
                            alt_text = ' '.join(text_matches)
                            print(f"   Alternative extraction: {len(alt_text)} chars")
                            print(f"   Preview: {alt_text[:200]}...")
                            return alt_text
                
                return ""
            
        except Exception as e:
            print(f"❌ Error with python-docx: {e}")
            return ""
            
    except Exception as e:
        print(f"❌ Error downloading .docx file: {e}")
        return ""

def test_openai_extraction(text):
    """Test OpenAI parameter extraction"""
    print(f"\n🤖 Testing OpenAI extraction on {len(text)} characters")
    print("=" * 50)
    
    if not text.strip():
        print("❌ No text to process")
        return {}
    
    try:
        # Initialize OpenAI client
        client = openai.OpenAI(api_key=config.OPENAI_API_KEY)
        
        prompt = """
Ты эксперт по анализу объявлений о продаже квартир в Израиле. 

Проанализируй следующий текст объявления и извлеки параметры в JSON формате:

1. "rooms" - количество комнат (например: "3", "4", "5", "2.5", "4.5")
2. "floor" - этаж (например: "2", "5", "10", "высокий этаж" -> "10+")  
3. "furniture" - мебель (если упоминается мебель/обставлена -> "да", иначе -> "нет")
4. "price" - цена в шекелях (только цифры, без валюты)
5. "district" - район Ришон-ле-Циона (ищи названия районов как "НАХЛАД ИУДА", "РЕМЕЗ", "НЕВЕ ДЕНЯ" и т.д.)

Если параметр не найден, не включай его в ответ.

Верни только JSON без дополнительного текста.

Текст объявления:
"""
        
        print("📤 Sending to OpenAI...")
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an expert at extracting apartment parameters from Russian real estate ads."},
                {"role": "user", "content": f"{prompt}\n\n{text}"}
            ],
            max_tokens=500,
            temperature=0.1
        )
        
        content = response.choices[0].message.content.strip()
        print(f"📥 OpenAI response: {content}")
        
        # Try to parse JSON
        import json
        try:
            parameters = json.loads(content)
            print(f"✅ Parsed parameters: {parameters}")
            return parameters
        except json.JSONDecodeError:
            print(f"⚠️ Could not parse as JSON, raw response: {content}")
            return {}
            
    except Exception as e:
        print(f"❌ OpenAI extraction failed: {e}")
        return {}

def main():
    print("🧪 DOCX EXTRACTION AND OPENAI TEST")
    print("=" * 50)
    
    # Setup API
    service = setup_drive_api()
    
    # Test with the specific file from the first folder
    file_id = "1OBQr4XDqdGkNQ1xHn9VSrpNjBz-3LM9I"  # бэн аришон.docx
    filename = "бэн аришон.docx"
    
    # Test .docx extraction
    extracted_text = test_docx_extraction(service, file_id, filename)
    
    # Test OpenAI extraction if we got text
    if extracted_text:
        parameters = test_openai_extraction(extracted_text)
        
        print(f"\n🎯 FINAL RESULTS:")
        print(f"   📄 Text length: {len(extracted_text)} characters")
        print(f"   🔧 Parameters: {parameters}")
    else:
        print("\n❌ Cannot test OpenAI - no text extracted")
    
    print("\n🏁 Test completed!")

if __name__ == "__main__":
    main() 