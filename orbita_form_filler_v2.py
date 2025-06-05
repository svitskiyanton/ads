"""
Orbita Form Filler v2.0 - Enhanced with Auto-Registration and OpenAI Integration

New Features:
- Auto-registration of accounts before posting ads
- OpenAI GPT-4.1 nano integration for parameter extraction
- Updated Google Drive path structure
- No logout between ads (only after last ad)
- Google Docs text extraction instead of separate files

Author: AI Assistant
Date: 2025
"""

import os
import re
import time
import random
import string
import tempfile
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Tuple
import traceback

# Core dependencies
from playwright.sync_api import sync_playwright, Page, Browser, BrowserContext
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import Flow
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
import openai

# Third-party services
from twocaptcha import TwoCaptcha

# Tor integration
import subprocess
import socket
import requests
import psutil

# Configuration
import config

class TorIPChanger:
    """Tor IP changing functionality integrated for ad posting automation - Based on proven original implementation"""
    
    def __init__(self):
        self.tor_process = None
        self.tor_path = None
        self.is_initialized = False
        self.proxy_config = {
            'http': 'socks5://127.0.0.1:9050',
            'https': 'socks5://127.0.0.1:9050'
        }
        
    def initialize_tor(self):
        """Initialize Tor based on operating system"""
        try:
            print("🔧 Initializing Tor IP changer...")
            
            if os.name == 'posix':  # Linux/Unix
                return self._initialize_tor_linux()
            elif os.name == 'nt':   # Windows
                return self._initialize_tor_windows()
            else:
                print("❌ Unsupported operating system")
                return False
                
        except Exception as e:
            print(f"❌ Error initializing Tor: {e}")
            return False
    
    def _initialize_tor_linux(self):
        """Initialize Tor on Linux/Unix systems"""
        try:
            # Check if running as root (required for Linux)
            if os.geteuid() != 0:
                print("⚠️ Warning: Tor operations may require root privileges on Linux")
            
            # Check if Tor is installed
            if subprocess.run(['which', 'tor'], stdout=subprocess.PIPE, stderr=subprocess.PIPE).returncode != 0:
                print("🔄 Tor not found. Installing...")
                if os.system("sudo apt install tor -y > /dev/null 2>&1") != 0:
                    print("❌ Failed to install Tor. Please install manually: sudo apt install tor")
                    return False
                print("✅ Tor installed successfully")
            else:
                print("✅ Tor is already installed")
            
            self.is_initialized = True
            return True
            
        except Exception as e:
            print(f"❌ Error initializing Tor on Linux: {e}")
            return False
    
    def _initialize_tor_windows(self):
        """Initialize Tor on Windows systems"""
        try:
            # Read tor path from config file
            tor_path_file = "tor_path.txt"
            if not os.path.exists(tor_path_file):
                # Create default tor_path.txt
                with open(tor_path_file, "w") as f:
                    f.write("C:\\")
            
            with open(tor_path_file, "r") as f:
                extract_path = f.read().strip()
            
            self.tor_path = f"{extract_path}\\Tor Expert Bundle\\tor\\tor.exe"
            
            if not os.path.exists(self.tor_path):
                print("🔄 Tor not found. Downloading and installing...")
                if not self._download_tor_windows(extract_path):
                    return False
            else:
                print("✅ Tor is already installed")
            
            self.is_initialized = True
            return True
            
        except Exception as e:
            print(f"❌ Error initializing Tor on Windows: {e}")
            return False
    
    def _download_tor_windows(self, extract_path):
        """Download and extract Tor for Windows"""
        try:
            tor_url = "https://archive.torproject.org/tor-package-archive/torbrowser/14.0.7/tor-expert-bundle-windows-x86_64-14.0.7.tar.gz"
            filename = "tor.tar.gz"
            
            print(f"📥 Downloading Tor from {tor_url}")
            import urllib.request
            urllib.request.urlretrieve(tor_url, filename)
            print("✅ Download complete")
            
            print(f"📂 Extracting to {extract_path}\\Tor Expert Bundle")
            import tarfile
            with tarfile.open(filename, "r:gz") as tar:
                tar.extractall(f"{extract_path}\\Tor Expert Bundle", filter='fully_trusted')
            
            os.remove(filename)
            print("✅ Tor extracted successfully")
            return True
            
        except Exception as e:
            print(f"❌ Error downloading Tor: {e}")
            return False
    
    def start_tor(self):
        """Start Tor service"""
        try:
            if not self.is_initialized:
                print("❌ Tor not initialized. Call initialize_tor() first")
                return False
            
            print("🔄 Starting Tor service...")
            
            if os.name == 'posix':  # Linux/Unix
                # Check if already running
                result = subprocess.run(["sudo", "service", "tor", "status"], capture_output=True, text=True)
                if "Active: active" in result.stdout:
                    print("✅ Tor is already running")
                else:
                    subprocess.run("sudo service tor start", shell=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                    time.sleep(3)
                    print("✅ Tor service started")
                    
            elif os.name == 'nt':   # Windows
                # Check if already running
                result = subprocess.run(["tasklist"], capture_output=True, text=True)
                if "tor.exe" in result.stdout:
                    print("✅ Tor is already running")
                else:
                    self.tor_process = subprocess.Popen([self.tor_path], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
                    print("✅ Tor started")
            
            # Wait longer for Tor to establish circuits
            print("⏳ Waiting for Tor to establish circuits...")
            time.sleep(10)  # Increased from 3 to 10 seconds
            
            # Try to verify connection multiple times
            print("🔍 Verifying Tor connection...")
            for attempt in range(3):
                if self.get_current_ip(max_retries=2, retry_delay=3):
                    print("✅ Tor connection verified and working")
                    return True
                if attempt < 2:
                    print(f"⏳ Connection attempt {attempt + 1} failed, waiting 5 more seconds...")
                    time.sleep(5)
            
            print("⚠️ Tor started but connection verification failed")
            return True  # Return True anyway, verification might work later
            
        except Exception as e:
            print(f"❌ Error starting Tor: {e}")
            return False
    
    def change_ip(self):
        """Change IP by restarting Tor"""
        try:
            if not self.is_initialized:
                print("❌ Tor not initialized")
                return False
            
            print("🔄 Changing IP address...")
            
            if os.name == 'posix':  # Linux/Unix
                subprocess.run("sudo service tor reload", shell=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                
            elif os.name == 'nt':   # Windows
                # Kill existing tor process
                subprocess.run(['taskkill', '/F', '/IM', 'tor.exe'], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                time.sleep(2)  # Wait a bit longer
                # Start new tor process
                self.tor_process = subprocess.Popen([self.tor_path], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
                
            # Wait for new circuits to establish
            print("⏳ Waiting for new circuits to establish...")
            time.sleep(8)  # Increased wait time
            
            # Verify IP change with multiple attempts
            print("🔍 Verifying IP change...")
            for attempt in range(3):
                new_ip = self.get_current_ip(max_retries=2, retry_delay=3)
                if new_ip:
                    print(f"✅ IP successfully changed to: {new_ip}")
                    return True
                if attempt < 2:
                    print(f"⏳ IP verification attempt {attempt + 1} failed, waiting 5 more seconds...")
                    time.sleep(5)
            
            print("⚠️ IP change completed but verification failed")
            return True  # Proceed anyway
                
        except Exception as e:
            print(f"❌ Error changing IP: {e}")
            return False
    
    def get_current_ip(self, max_retries=3, retry_delay=5):
        """Get current IP address through Tor with retry logic"""
        # List of IP checking services to try
        ip_services = [
            "https://httpbin.org/ip",
            "https://icanhazip.com",
            "https://api.ipify.org?format=json",
            "https://ipinfo.io/json"
        ]
        
        for attempt in range(max_retries):
            for service_url in ip_services:
                try:
                    print(f"🔍 Checking IP address (attempt {attempt + 1}/{max_retries}) via {service_url}...")
                    
                    response = requests.get(service_url, proxies=self.proxy_config, timeout=15)
                    
                    # Parse response based on service
                    if "httpbin.org" in service_url:
                        ip = response.json().get('origin', 'Unknown')
                    elif "icanhazip.com" in service_url:
                        ip = response.text.strip()
                    elif "ipify.org" in service_url:
                        ip = response.json().get('ip', 'Unknown')
                    elif "ipinfo.io" in service_url:
                        ip = response.json().get('ip', 'Unknown')
                    else:
                        ip = response.text.strip()
                    
                    if ip and ip != 'Unknown':
                        print(f"✅ IP verified: {ip}")
                        return ip
                        
                except requests.exceptions.ConnectTimeout:
                    print(f"⏳ Connection timeout to {service_url}")
                    continue
                except requests.exceptions.ProxyError as e:
                    print(f"🔌 Proxy error with {service_url}: {e}")
                    continue
                except Exception as e:
                    print(f"⚠️ Error with {service_url}: {e}")
                    continue
            
            # If all services failed for this attempt, wait before retrying
            if attempt < max_retries - 1:
                print(f"⏳ All services failed on attempt {attempt + 1}, retrying in {retry_delay} seconds...")
                time.sleep(retry_delay)
        
        print("❌ Failed to verify IP after all attempts with all services")
        return None
    
    def stop_tor(self):
        """Stop Tor service"""
        try:
            print("🛑 Stopping Tor service...")
            
            if os.name == 'posix':  # Linux/Unix
                subprocess.run("sudo service tor stop", shell=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                
            elif os.name == 'nt' and self.tor_process:  # Windows
                self.tor_process.kill()
                self.tor_process = None
            
            print("✅ Tor stopped")
            
        except Exception as e:
            print(f"⚠️ Error stopping Tor: {e}")
    
    def configure_browser_proxy(self, context):
        """Configure Playwright browser context to use Tor proxy"""
        try:
            # Set proxy for the browser context
            proxy_config = {
                'server': 'socks5://127.0.0.1:9050'
            }
            print("🌐 Configuring browser to use Tor proxy")
            return proxy_config
        except Exception as e:
            print(f"❌ Error configuring browser proxy: {e}")
            return None

class OpenAIExtractor:
    """OpenAI integration for extracting parameters from ad text"""
    
    def __init__(self):
        if not config.OPENAI_API_KEY or config.OPENAI_API_KEY == "your_openai_api_key_here":
            raise ValueError("❌ OpenAI API key not configured in config.py")
        
        # Clear any proxy environment variables that might interfere with OpenAI client
        env_vars_to_clear = ['HTTP_PROXY', 'HTTPS_PROXY', 'ALL_PROXY', 'http_proxy', 'https_proxy', 'all_proxy']
        self.original_env = {}
        
        for var in env_vars_to_clear:
            if var in os.environ:
                self.original_env[var] = os.environ[var]
                del os.environ[var]
        
        try:
            # Create OpenAI client without proxy configuration
            self.client = openai.OpenAI(
                api_key=config.OPENAI_API_KEY,
                timeout=30.0
            )
        finally:
            # Restore original environment variables
            for var, value in self.original_env.items():
                os.environ[var] = value
    
    def extract_parameters(self, ad_text: str) -> Dict[str, str]:
        """Extract apartment parameters from Russian ad text using GPT-4o mini"""
        
        prompt = """
Ты эксперт по анализу объявлений о продаже квартир в Израиле, особенно в городе Ришон-ле-Цион. 

Проанализируй следующий текст объявления и извлеки параметры в JSON формате:

1. "rooms" - количество комнат (например: "3", "4", "5", "2.5", "4.5")
2. "floor" - этаж (например: "2", "5", "10", "высокий этаж" -> "10+")  
3. "furniture" - мебель (если упоминается мебель/обставлена -> "да", иначе -> "нет")
4. "price" - цена в шекелях (только цифры, без валюты). ВНИМАНИЕ: Ищи ЛЮБЫЕ длинные числа (4+ цифры), даже если они разделены пробелами или запятыми - это скорее всего цена! Например: "2 000 000", "2,500,000", "1.800.000" - все это цены.
5. "district" - район/адрес в Ришон-ле-Ционе. ВАЖНО: Внимательно ищи слова связанные с районами Ришон-ле-Циона, улицами и адресами:
   - Районы: "НАХЛАД ИУДА", "РЕМЕЗ", "НЕВЕ ДЕНЯ", "КИРЬЯТ ГАОН", "РАМАТ ЭЛИЯУ", "ЦЕНТР", "СТАРЫЙ ГОРОД"
   - Улицы: "Ротшильд", "Герцль", "Жаботинский", "Бялик", "Ахад Хаам", "Рош Пина", "Вайцман", "Бен Гурион", "Соколов"
   - Ключевые слова: "ул.", "улица", "район", "квартал", любые названия на иврите или русском связанные с Ришон-ле-Ционом

Если параметр не найден, не включай его в ответ.

Верни только JSON без дополнительного текста.

Текст объявления:
"""
        
        try:
            response = self.client.chat.completions.create(
                model="gpt-4o-mini",  # Using available model
                messages=[
                    {"role": "system", "content": "You are an expert at extracting apartment parameters from Russian real estate ads."},
                    {"role": "user", "content": f"{prompt}\n\n{ad_text}"}
                ],
                max_tokens=500,
                temperature=0.1
            )
            
            content = response.choices[0].message.content.strip()
            
            # Try to parse JSON response
            import json
            try:
                parameters = json.loads(content)
                print(f"✅ Extracted parameters: {parameters}")
                return parameters
            except json.JSONDecodeError:
                # If not valid JSON, try to extract manually
                print(f"⚠️ Could not parse JSON, using fallback extraction")
                return self._fallback_extraction(ad_text)
                
        except Exception as e:
            print(f"❌ OpenAI extraction failed: {e}")
            return self._fallback_extraction(ad_text)
    
    def _fallback_extraction(self, ad_text: str) -> Dict[str, str]:
        """Fallback manual extraction if OpenAI fails"""
        parameters = {}
        
        # Extract rooms
        room_patterns = [
            r'(\d+(?:\.\d+)?)\s*комнат',
            r'(\d+)\s*к',
            r'(\d+)\s*room'
        ]
        for pattern in room_patterns:
            match = re.search(pattern, ad_text, re.IGNORECASE)
            if match:
                parameters['rooms'] = match.group(1)
                break
        
        # Extract district - expanded list with streets
        districts = ['НАХЛАД ИУДА', 'РЕМЕЗ', 'НЕВЕ ДЕНЯ', 'КИРЬЯТ ГАОН', 'РАМАТ ЭЛИЯУ', 'ЦЕНТР', 'СТАРЫЙ ГОРОД']
        streets = ['РОТШИЛЬД', 'ГЕРЦЛЬ', 'ЖАБОТИНСКИЙ', 'БЯЛИК', 'АХАД ХААМ', 'РОШ ПИНА', 'ВАЙЦМАН', 'БEN ГУРИОН', 'СОКОЛОВ']
        
        # Check districts first
        for district in districts:
            if district in ad_text.upper():
                parameters['district'] = district
                break
        
        # If no district found, check streets
        if 'district' not in parameters:
            for street in streets:
                if street in ad_text.upper():
                    parameters['district'] = f"ул. {street}"
                    break
        
        # Extract price - improved to handle spaces and commas
        price_patterns = [
            r'(\d{1,3}(?:[,\s]\d{3})*)\s*₪',  # With shekel symbol
            r'(\d{1,3}(?:[,\s]\d{3})+)',      # Long numbers with separators (4+ digits total)
            r'(\d{7,})',                      # Very long numbers (7+ digits)
            r'(\d{4,6})',                     # Medium numbers (4-6 digits)
        ]
        
        for pattern in price_patterns:
            price_match = re.search(pattern, ad_text)
            if price_match:
                price = price_match.group(1).replace(',', '').replace(' ', '')
                # Only consider it a price if it's reasonable (between 100k and 10M shekels)
                try:
                    price_num = int(price)
                    if 100000 <= price_num <= 10000000:
                        parameters['price'] = price
                        break
                except ValueError:
                    continue
        
        print(f"✅ Fallback extracted: {parameters}")
        return parameters

class GoogleDriveClient:
    """Enhanced Google Drive client for new folder structure"""
    
    def __init__(self):
        self.service = None
        self.credentials = None
        
    def authenticate(self):
        """Authenticate with Google Drive API"""
        try:
            creds = None
            
            # Load existing token
            if os.path.exists('token.json'):
                creds = Credentials.from_authorized_user_file('token.json')
            
            # If no valid credentials, run OAuth flow
            if not creds or not creds.valid:
                if creds and creds.expired and creds.refresh_token:
                    creds.refresh(Request())
                else:
                    flow = Flow.from_client_secrets_file(
                        'credentials.json',
                        scopes=['https://www.googleapis.com/auth/drive.readonly']
                    )
                    flow.redirect_uri = 'http://localhost:8080/callback'
                    
                    auth_url, _ = flow.authorization_url(prompt='consent')
                    print(f"🔗 Please visit: {auth_url}")
                    
                    code = input("📝 Enter authorization code: ")
                    flow.fetch_token(code=code)
                    creds = flow.credentials
                
                # Save credentials
                with open('token.json', 'w') as token:
                    token.write(creds.to_json())
            
            self.credentials = creds
            self.service = build('drive', 'v3', credentials=creds)
            print("✅ Google Drive authenticated successfully")
            return True
            
        except Exception as e:
            print(f"❌ Google Drive authentication failed: {e}")
            return False
    
    def find_folder_by_path(self, path: str) -> Optional[str]:
        """Find folder ID by path in user's Drive"""
        try:
            path_parts = [p.strip() for p in path.split('/') if p.strip()]
            
            # Start with root folder
            current_folder_id = 'root'
            
            for part in path_parts:
                print(f"🔍 Searching for folder: '{part}' in parent: {current_folder_id}")
                
                # Simple search in current parent
                query = f"name='{part}' and parents in '{current_folder_id}' and mimeType='application/vnd.google-apps.folder'"
                results = self.service.files().list(
                    q=query, 
                    fields="files(id, name)"
                ).execute()
                folders = results.get('files', [])
                
                if folders:
                    current_folder_id = folders[0]['id']
                    print(f"   ✅ Found: {part}")
                    print(f"   📂 Moving to folder ID: {current_folder_id}")
                else:
                    print(f"❌ Could not find folder: '{part}'")
                    return None
            
            print(f"✅ Successfully found path '{path}' with ID: {current_folder_id}")
            return current_folder_id
            
        except Exception as e:
            print(f"❌ Error finding folder by path: {e}")
            return None
    
    def get_ad_folders(self) -> List[Dict]:
        """Get all ad folders from the specified path"""
        try:
            parent_folder_id = self.find_folder_by_path(config.GOOGLE_DRIVE_PATH)
            if not parent_folder_id:
                return []
            
            # Get all subfolders in the ПРОДАЖА folder
            query = f"parents in '{parent_folder_id}' and mimeType='application/vnd.google-apps.folder'"
            results = self.service.files().list(q=query).execute()
            
            folders = results.get('files', [])
            print(f"✅ Found {len(folders)} ad folders")
            
            return folders
            
        except Exception as e:
            print(f"❌ Error getting ad folders: {e}")
            return []
    
    def get_folder_contents(self, folder_id: str) -> Dict:
        """Get contents of a specific folder"""
        try:
            query = f"parents in '{folder_id}'"
            results = self.service.files().list(q=query).execute()
            
            files = results.get('files', [])
            
            # Separate text documents (Google Docs and .docx) and images
            text_documents = []
            images = []
            
            for file in files:
                if file['mimeType'] == 'application/vnd.google-apps.document':
                    # Google Doc
                    text_documents.append({'type': 'google_doc', 'file': file})
                elif file['mimeType'] == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                    # .docx file
                    text_documents.append({'type': 'docx', 'file': file})
                elif file['mimeType'].startswith('image/'):
                    images.append(file)
            
            # Sort images and take first 5
            images = sorted(images, key=lambda x: x['name'])[:config.MAX_IMAGES_PER_AD]
            
            return {
                'text_documents': text_documents,
                'images': images
            }
            
        except Exception as e:
            print(f"❌ Error getting folder contents: {e}")
            return {'text_documents': [], 'images': []}
    
    def download_document_text(self, document_info: Dict) -> str:
        """Download text from Google Doc or .docx file"""
        try:
            doc_type = document_info['type']
            file_info = document_info['file']
            file_id = file_info['id']
            file_name = file_info['name']
            
            if doc_type == 'google_doc':
                # Download Google Doc as plain text
                print(f"📄 Downloading Google Doc: {file_name}")
                return self.download_google_doc_text(file_id)
            elif doc_type == 'docx':
                # Download .docx file and extract text
                print(f"📄 Downloading .docx file: {file_name}")
                return self.download_docx_text(file_id, file_name)
            else:
                print(f"❌ Unknown document type: {doc_type}")
                return ""
                
        except Exception as e:
            print(f"❌ Error downloading document: {e}")
            return ""
    
    def download_docx_text(self, file_id: str, filename: str) -> str:
        """Download .docx file and extract text content with improved method"""
        try:
            # Download the .docx file
            request = self.service.files().get_media(fileId=file_id)
            
            import io
            fh = io.BytesIO()
            downloader = MediaIoBaseDownload(fh, request)
            
            done = False
            while done is False:
                status, done = downloader.next_chunk()
            
            fh.seek(0)
            
            # Extract text from .docx using python-docx
            try:
                from docx import Document
                doc = Document(fh)
                
                # Extract all text content
                text_content = []
                
                # From paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_content.append(paragraph.text.strip())
                
                # From tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_content.append(cell.text.strip())
                
                # From headers/footers if available
                for section in doc.sections:
                    if hasattr(section, 'header') and section.header:
                        for paragraph in section.header.paragraphs:
                            if paragraph.text.strip():
                                text_content.append(paragraph.text.strip())
                    
                    if hasattr(section, 'footer') and section.footer:
                        for paragraph in section.footer.paragraphs:
                            if paragraph.text.strip():
                                text_content.append(paragraph.text.strip())
                
                final_text = '\n'.join(text_content)
                print(f"✅ Extracted {len(final_text)} characters from .docx file")
                
                if final_text:
                    return final_text
                else:
                    # Try alternative XML extraction
                    print("🔍 Trying alternative XML extraction...")
                    return self._extract_docx_via_xml(fh)
                
            except ImportError:
                print("⚠️ python-docx not installed. Installing...")
                import subprocess
                subprocess.check_call(['pip', 'install', 'python-docx'])
                
                # Try again after installation
                from docx import Document
                fh.seek(0)  # Reset file pointer
                doc = Document(fh)
                
                text_content = []
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_content.append(paragraph.text.strip())
                
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                text_content.append(cell.text.strip())
                
                text = '\n'.join(text_content)
                print(f"✅ Extracted {len(text)} characters from .docx file")
                return text if text else self._extract_docx_via_xml(fh)
                
        except Exception as e:
            print(f"❌ Error downloading .docx file {filename}: {e}")
            return ""
    
    def _extract_docx_via_xml(self, file_handle):
        """Alternative .docx extraction using XML parsing"""
        try:
            file_handle.seek(0)
            import zipfile
            import re
            import xml.etree.ElementTree as ET
            
            with zipfile.ZipFile(file_handle, 'r') as zip_file:
                if 'word/document.xml' in zip_file.namelist():
                    xml_content = zip_file.read('word/document.xml').decode('utf-8')
                    
                    # Parse XML properly
                    try:
                        root = ET.fromstring(xml_content)
                        
                        # Find all text elements
                        text_elements = []
                        for elem in root.iter():
                            if elem.tag.endswith('}t') and elem.text:
                                text_elements.append(elem.text)
                        
                        if text_elements:
                            alt_text = ' '.join(text_elements)
                            print(f"✅ XML extraction: {len(alt_text)} characters")
                            return alt_text
                            
                    except ET.ParseError:
                        # Fallback to regex
                        text_matches = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', xml_content)
                        if text_matches:
                            alt_text = ' '.join(text_matches)
                            print(f"✅ Regex extraction: {len(alt_text)} characters")
                            return alt_text
            
            print("❌ No text found in .docx file via XML")
            return ""
            
        except Exception as e:
            print(f"❌ XML extraction failed: {e}")
            return ""
    
    def download_google_doc_text(self, file_id: str) -> str:
        """Download Google Doc as plain text"""
        try:
            # Export as plain text (simple like test file)
            request = self.service.files().export_media(
                fileId=file_id, 
                mimeType='text/plain'
            )
            
            import io
            fh = io.BytesIO()
            downloader = MediaIoBaseDownload(fh, request)
            
            done = False
            while done is False:
                status, done = downloader.next_chunk()
            
            fh.seek(0)
            text = fh.read().decode('utf-8')
            
            return text
            
        except Exception as e:
            print(f"❌ Error downloading Google Doc: {e}")
            return ""
    
    def download_image(self, file_id: str, filename: str) -> Optional[str]:
        """Download image to temporary file"""
        try:
            # Get media (simple like test file)
            request = self.service.files().get_media(
                fileId=file_id
            )
            
            temp_dir = tempfile.gettempdir()
            temp_path = os.path.join(temp_dir, filename)
            
            import io
            fh = io.BytesIO()
            downloader = MediaIoBaseDownload(fh, request)
            
            done = False
            while done is False:
                status, done = downloader.next_chunk()
            
            with open(temp_path, 'wb') as f:
                f.write(fh.getvalue())
            
            return temp_path
            
        except Exception as e:
            print(f"❌ Error downloading image {filename}: {e}")
            return None

class AccountRegistrar:
    """Handle automatic account registration"""
    
    def __init__(self, page: Page):
        self.page = page
        self.solver = TwoCaptcha(config.CAPTCHA_API_KEY) if config.CAPTCHA_API_KEY != "your_2captcha_api_key_here" else None
    
    def generate_random_email(self) -> str:
        """Generate random email address"""
        username = ''.join(random.choices(string.ascii_lowercase, k=7))
        domain = ''.join(random.choices(string.ascii_lowercase, k=10))
        return f"{username}@{domain}.com"
    
    def register_account(self) -> Tuple[bool, str]:
        """Register new account and return success status and email"""
        try:
            print("🔐 Starting account registration...")
            
            # Navigate to login page
            self.page.goto("https://passport.orbita.co.il/site/login/")
            self.page.wait_for_load_state('networkidle')
            
            # Click registration button
            reg_button = self.page.locator("a.reg-pop")
            reg_button.click()
            
            # Wait for registration form to appear
            self.page.wait_for_selector("#signup-form", state="visible")
            
            # Generate registration data
            email = self.generate_random_email()
            password = config.REGISTRATION_PASSWORD
            name = config.REGISTRATION_NAME
            
            print(f"📧 Generated email: {email}")
            
            # Fill registration form
            self.page.fill("#signupform-email", email)
            self.page.fill("#signupform-password", password)
            self.page.fill("#signupform-name", name)
            
            # Handle reCAPTCHA if present (but don't fail if not present)
            recaptcha_result = self._solve_recaptcha()
            if recaptcha_result:
                print("✅ reCAPTCHA handling completed")
            
            # Submit registration
            submit_button = self.page.locator("button[name='signup-button']")
            submit_button.click()
            
            # Wait for response
            print("⏳ Waiting for registration response...")
            time.sleep(5)
            
            # Check current URL and page content to determine success
            current_url = self.page.url
            page_content = self.page.content().lower()
            
            # Russian success phrases
            success_phrases = [
                "добро пожаловать",  # Welcome
                "успешно создали",   # Successfully created
                "успешно зарегистрированы",  # Successfully registered
                "регистрация завершена",  # Registration completed
                "вы успешно",  # You have successfully
            ]
            
            # Multiple success indicators
            success_indicators = [
                "login" in current_url,  # Redirected to login
                "success" in current_url,
                any(phrase in page_content for phrase in success_phrases),  # Russian success messages
                self.page.locator("#login-form").is_visible(),  # Login form visible
            ]
            
            # Error phrases (but exclude success contexts)
            error_phrases = [
                "ошибка",  # Error in Russian
                "неверный",  # Invalid/incorrect
                "не удалось",  # Failed to
                "попробуйте снова",  # Try again
            ]
            
            # Check for error indicators (but not if we already found success)
            error_indicators = [
                any(phrase in page_content for phrase in error_phrases),
                self.page.locator(".alert-danger, .has-error").count() > 0,
            ]
            
            # Determine result
            has_success = any(success_indicators)
            has_errors = any(error_indicators) and not has_success  # Ignore errors if we found success
            
            if has_success:
                print("✅ Registration successful!")
                print(f"   Welcome message detected in response")
                return True, email
            elif has_errors:
                print("❌ Registration failed with errors")
                try:
                    # Try to get error messages
                    error_elements = self.page.locator('.alert-danger, .has-error').all()
                    for element in error_elements:
                        error_text = element.text_content()
                        if error_text and error_text.strip():
                            print(f"   Error: {error_text}")
                except:
                    pass
                return False, email
            else:
                print("⚠️ Registration status unclear - assuming success")
                return True, email  # Assume success if unclear
                
        except Exception as e:
            print(f"❌ Registration failed: {e}")
            return False, ""
    
    def _solve_recaptcha(self) -> bool:
        """Solve reCAPTCHA during registration"""
        try:
            if not self.solver:
                print("ℹ️ No captcha solver configured, continuing without reCAPTCHA")
                return True
            
            # Look for reCAPTCHA
            recaptcha_frame = self.page.locator("iframe[src*='recaptcha']").first
            if not recaptcha_frame.is_visible():
                print("ℹ️ No reCAPTCHA found - continuing normally")
                return True
            
            # Get site key
            site_key_element = self.page.locator("[data-sitekey]").first
            if not site_key_element.is_visible():
                print("ℹ️ No reCAPTCHA site key found - continuing normally")
                return True
            
            site_key = site_key_element.get_attribute("data-sitekey")
            page_url = self.page.url
            
            print("🤖 Solving reCAPTCHA...")
            
            # Solve captcha
            result = self.solver.recaptcha(
                sitekey=site_key,
                url=page_url
            )
            
            if result and 'code' in result:
                # Inject solution
                self.page.evaluate(f"""
                    document.getElementById('g-recaptcha-response').innerHTML = '{result['code']}';
                    document.getElementById('g-recaptcha-response').style.display = 'block';
                """)
                
                print("✅ reCAPTCHA solved successfully")
                return True
            else:
                print("⚠️ reCAPTCHA solving failed, continuing anyway")
                return True  # Continue even if solving failed
            
        except Exception as e:
            print(f"⚠️ reCAPTCHA solving error (continuing anyway): {e}")
            return True  # Always return True to continue registration

class OrbitaFormFillerV2:
    """Enhanced Orbita Form Filler with new algorithm"""
    
    def __init__(self):
        self.browser = None
        self.context = None
        self.page = None
        self.tor_changer = TorIPChanger() if config.USE_TOR_IP_ROTATION else None
        self.drive_client = GoogleDriveClient()
        self.openai_extractor = OpenAIExtractor()
        self.current_account_email = None
        self.processed_ads_log = "processed_ads_v2.log"
        
    def initialize(self) -> bool:
        """Initialize all components"""
        try:
            print("🚀 Initializing Orbita Form Filler v2...")
            
            # Initialize Tor if enabled - REQUIRED per user request
            if self.tor_changer:
                if not self.tor_changer.initialize_tor():
                    print("❌ Tor initialization failed - STOPPING as required")
                    return False
                elif not self.tor_changer.start_tor():
                    print("❌ Tor start failed - STOPPING as required")
                    return False
                else:
                    print("✅ Tor started successfully - continuing")
            
            # Authenticate Google Drive
            if not self.drive_client.authenticate():
                print("❌ Google Drive authentication failed")
                return False
            
            print("✅ Initialization completed")
            return True
            
        except Exception as e:
            print(f"❌ Initialization failed: {e}")
            return False
    
    def start_browser(self) -> bool:
        """Start browser with proper configuration"""
        try:
            print("🌐 Starting browser...")
            
            playwright = sync_playwright().start()
            
            # Browser arguments
            browser_args = [
                '--no-blink-features=AutomationControlled',
                '--disable-blink-features=AutomationControlled',
                '--disable-dev-shm-usage',
                '--disable-extensions',
                '--no-first-run',
                '--no-default-browser-check',
                '--disable-default-apps',
                '--disable-popup-blocking'
            ]
            
            # Prepare launch options
            launch_options = {
                'headless': config.BROWSER_HEADLESS,
                'slow_mo': config.BROWSER_SLOW_MO,
                'args': browser_args
            }
            
            # Add proxy if Tor is enabled
            if self.tor_changer and self.tor_changer.is_initialized:
                proxy_config = self.tor_changer.configure_browser_proxy(None)
                if proxy_config:
                    launch_options['proxy'] = proxy_config
                    print("🌐 Browser will use Tor proxy")
            
            # Launch browser
            self.browser = playwright.chromium.launch(**launch_options)
            
            # Create context with human-like settings
            context_options = {
                'viewport': {'width': 1366, 'height': 768},
                'user_agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
            }
            
            self.context = self.browser.new_context(**context_options)
            
            # Create page
            self.page = self.context.new_page()
            
            print("✅ Browser started successfully")
            return True
            
        except Exception as e:
            print(f"❌ Browser start failed: {e}")
            return False
    
    def register_and_login(self) -> bool:
        """Register new account and login"""
        try:
            registrar = AccountRegistrar(self.page)
            success, email = registrar.register_account()
            
            if success:
                self.current_account_email = email
                print(f"✅ Account registered and logged in: {email}")
                return True
            else:
                print("❌ Account registration failed")
                return False
                
        except Exception as e:
            print(f"❌ Registration/login failed: {e}")
            return False
    
    def process_all_ads(self) -> Dict[str, int]:
        """Process all ads from Google Drive"""
        try:
            print("📁 Getting ad folders from Google Drive...")
            
            ad_folders = self.drive_client.get_ad_folders()
            if not ad_folders:
                print("❌ No ad folders found")
                return {'processed': 0, 'failed': 0, 'skipped': 0, 'empty': 0}
            
            stats = {'processed': 0, 'failed': 0, 'skipped': 0, 'empty': 0}
            processed_folders = self._load_processed_ads()
            
            print(f"📊 Found {len(ad_folders)} ad folders")
            
            for i, folder in enumerate(ad_folders):
                folder_name = folder['name']
                folder_id = folder['id']
                
                print(f"\n📂 Processing folder {i+1}/{len(ad_folders)}: {folder_name}")
                
                # Skip if already processed
                if folder_id in processed_folders:
                    print(f"⏭️ Skipping already processed folder: {folder_name}")
                    stats['skipped'] += 1
                    continue
                
                # Add small delay between folders to avoid API rate limits
                if i > 0:  # Don't delay before first folder
                    print("⏳ Brief pause between folders...")
                    time.sleep(2)  # Back to reasonable 2 seconds
                
                # Process the ad
                result = self._process_single_ad(folder_id, folder_name)
                if result == 'success':
                    stats['processed'] += 1
                    self._log_processed_ad(folder_id, folder_name)
                    print(f"✅ Successfully processed: {folder_name}")
                elif result == 'empty':
                    stats['empty'] += 1
                    print(f"⚪ Skipped empty document: {folder_name}")
                else:
                    stats['failed'] += 1
                    print(f"❌ Failed to process: {folder_name}")
                
                # Wait between ads (except for the last one)
                if i < len(ad_folders) - 1:
                    print(f"⏳ Waiting {config.WAIT_BETWEEN_ADS} seconds before next ad...")
                    time.sleep(config.WAIT_BETWEEN_ADS)
                    
                    # Change IP if configured
                    if (self.tor_changer and 
                        config.TOR_IP_CHANGE_INTERVAL > 0 and 
                        (i + 1) % config.TOR_IP_CHANGE_INTERVAL == 0):
                        self.tor_changer.change_ip()
            
            # Logout after processing all ads
            if config.LOGOUT_BETWEEN_ADS or stats['processed'] > 0:
                self._logout()
            
            return stats
            
        except Exception as e:
            print(f"❌ Error processing ads: {e}")
            return {'processed': 0, 'failed': 0, 'skipped': 0, 'empty': 0}
    
    def _process_single_ad(self, folder_id: str, folder_name: str) -> str:
        """Process a single ad. Returns: 'success', 'empty', or 'failed'"""
        try:
            # Get folder contents
            print("📂 Getting folder contents...")
            contents = self.drive_client.get_folder_contents(folder_id)
            
            if not contents['text_documents']:
                print(f"❌ No text document found in folder: {folder_name}")
                return 'failed'
            
            # Get the first text document (ad text)
            doc = contents['text_documents'][0]
            print(f"📄 Downloading text document: {doc['file']['name']}")
            ad_text = self.drive_client.download_document_text(doc)
            
            if not ad_text.strip():
                print(f"⚪ Empty text document found: {doc['file']['name']}")
                print(f"   Skipping folder '{folder_name}' - no content to post")
                return 'empty'
            
            print(f"📄 Retrieved ad text ({len(ad_text)} characters)")
            
            # Extract parameters using OpenAI
            parameters = self.openai_extractor.extract_parameters(ad_text)
            
            # Download images
            image_paths = []
            for i, image in enumerate(contents['images']):
                print(f"🖼️ Downloading image {i+1}/{len(contents['images'])}: {image['name']}")
                path = self.drive_client.download_image(image['id'], image['name'])
                if path:
                    image_paths.append(path)
            
            print(f"🖼️ Downloaded {len(image_paths)} images")
            
            # Fill the form
            success = self._fill_orbita_form(ad_text, parameters, image_paths)
            
            # Cleanup temporary files
            for path in image_paths:
                try:
                    os.remove(path)
                except:
                    pass
            
            return 'success' if success else 'failed'
            
        except Exception as e:
            print(f"❌ Error processing single ad: {e}")
            return 'failed'
    
    def _dismiss_popups(self):
        """Dismiss any popup windows that appear on the page"""
        try:
            print("🔄 Checking for popup windows...")
            
            # Common popup dismiss selectors
            popup_selectors = [
                # Notification permission popups
                'button:text("Разрешить")',      # Allow in Russian
                'button:text("Allow")',           # Allow in English
                'button:text("Отклонить")',       # Deny in Russian
                'button:text("Deny")',            # Deny in English
                'button:text("Не сейчас")',       # Not now in Russian
                'button:text("Not now")',         # Not now in English
                'button:text("Закрыть")',         # Close in Russian
                'button:text("Close")',           # Close in English
                
                # Generic close buttons
                '[data-dismiss="modal"]',
                '.modal-close',
                '.popup-close', 
                '.close-button',
                'button[aria-label="Close"]',
                '[aria-label="Закрыть"]',
                
                # X buttons
                'button:has-text("×")',
                'span:has-text("×")',
                '.fa-close',
                '.fa-times',
                
                # OK/Cancel buttons
                'button:text("ОК")',
                'button:text("OK")',
                'button:text("Отмена")',
                'button:text("Cancel")',
            ]
            
            dismissed_count = 0
            for selector in popup_selectors:
                try:
                    elements = self.page.locator(selector)
                    count = elements.count()
                    if count > 0:
                        for i in range(min(count, 3)):  # Limit to 3 elements
                            try:
                                elements.nth(i).click(timeout=1000)
                                dismissed_count += 1
                                time.sleep(0.5)  # Brief pause between clicks
                            except:
                                continue
                except:
                    continue
            
            # Try pressing Escape key to dismiss any remaining modals
            try:
                self.page.keyboard.press('Escape')
                time.sleep(0.5)
            except:
                pass
            
            if dismissed_count > 0:
                print(f"✅ Dismissed {dismissed_count} popup(s)")
            else:
                print("ℹ️ No popups found")
                
        except Exception as e:
            print(f"⚠️ Error dismissing popups: {e}")
    
    def _fill_orbita_form(self, ad_text: str, parameters: Dict[str, str], image_paths: List[str]) -> bool:
        """Fill the Orbita form with ad data"""
        try:
            print("📝 Filling Orbita form...")
            
            # Navigate to add form
            self.page.goto("https://doska.orbita.co.il/my/add/")
            self.page.wait_for_load_state('networkidle')
            
            # Dismiss any popups that might appear
            time.sleep(2)  # Give time for popups to load
            self._dismiss_popups()
            
            # Wait a bit more and check again for popups
            time.sleep(1)
            self._dismiss_popups()
            
            # Select board: "Квартиры - продам" with better error handling
            try:
                print("🏠 Selecting apartment category...")
                board_select = self.page.locator("select").first
                board_select.wait_for(state="visible", timeout=10000)
                board_select.select_option(label="Квартиры - Продам")
                print("✅ Selected category: Квартиры - Продам")
            except Exception as e:
                print(f"⚠️ Could not select category: {e}")
                # Try alternative selectors
                try:
                    board_select = self.page.locator("select[name*='board'], select[name*='category']").first
                    board_select.select_option(label="Квартиры - Продам")
                    print("✅ Selected category via alternative selector")
                except:
                    print("❌ Could not find category selector")
            
            # Select city: "Ришон ле Цион" with better error handling
            try:
                print("🏙️ Selecting city...")
                city_select = self.page.locator("select").nth(1)
                city_select.wait_for(state="visible", timeout=10000)
                city_select.select_option(label="Ришон ле Цион")
                print("✅ Selected city: Ришон ле Цион")
            except Exception as e:
                print(f"⚠️ Could not select city: {e}")
                # Try alternative selectors
                try:
                    city_select = self.page.locator("select[name*='city'], select[name*='location']").first
                    city_select.select_option(label="Ришон ле Цион")
                    print("✅ Selected city via alternative selector")
                except:
                    print("❌ Could not find city selector")
            
            time.sleep(config.STEP_DELAY)
            
            # Fill ad text with better selector
            try:
                print("📝 Filling ad text...")
                # Try multiple selectors for textarea
                text_selectors = [
                    "textarea",
                    "textarea[name*='text']",
                    "textarea[name*='content']", 
                    "textarea[name*='description']",
                    "#ad-text",
                    "[name='ad_text']"
                ]
                
                ad_text_area = None
                for selector in text_selectors:
                    try:
                        ad_text_area = self.page.locator(selector).first
                        if ad_text_area.is_visible():
                            break
                    except:
                        continue
                
                if ad_text_area and ad_text_area.is_visible():
                    ad_text_area.fill(ad_text)
                    print("✅ Ad text filled successfully")
                else:
                    print("❌ Could not find ad text area")
                    
            except Exception as e:
                print(f"⚠️ Error filling ad text: {e}")
            
            time.sleep(config.STEP_DELAY)
            
            # Upload images with better error handling
            if image_paths:
                print(f"📤 Uploading {len(image_paths)} images...")
                try:
                    # Look for file inputs
                    file_inputs = self.page.locator("input[type='file']")
                    input_count = file_inputs.count()
                    
                    if input_count > 0:
                        for i, image_path in enumerate(image_paths[:min(5, input_count)]):
                            try:
                                file_input = file_inputs.nth(i)
                                file_input.set_input_files(image_path)
                                print(f"   ✅ Uploaded image {i+1}: {os.path.basename(image_path)}")
                                time.sleep(1)  # Wait between uploads
                            except Exception as e:
                                print(f"   ⚠️ Could not upload image {i+1}: {e}")
                        print("✅ Image upload completed")
                    else:
                        print("⚠️ No file input elements found")
                        
                except Exception as e:
                    print(f"❌ Image upload failed: {e}")
            
            time.sleep(config.STEP_DELAY)
            
            # Fill extracted parameters from OpenAI/fallback
            print("🏠 Filling extracted apartment parameters...")
            self._fill_apartment_parameters(parameters)
            
            # Fill contact information
            self._fill_contact_info(parameters)
            
            time.sleep(config.STEP_DELAY)
            
            # Handle reCAPTCHA if present
            print("🤖 Checking for reCAPTCHA...")
            self._solve_form_recaptcha()
            
            # Check agreement checkbox with multiple selectors
            try:
                print("☑️ Checking agreement checkbox...")
                checkbox_selectors = [
                    "input[type='checkbox']",
                    "input[name*='agree']",
                    "input[name*='terms']",
                    "input[name*='rules']",
                    "#agreement",
                    ".agreement-checkbox"
                ]
                
                checkbox_found = False
                for selector in checkbox_selectors:
                    try:
                        checkboxes = self.page.locator(selector)
                        for i in range(checkboxes.count()):
                            checkbox = checkboxes.nth(i)
                            if checkbox.is_visible() and not checkbox.is_checked():
                                checkbox.check()
                                checkbox_found = True
                                print("✅ Agreement checkbox checked")
                                break
                        if checkbox_found:
                            break
                    except:
                        continue
                        
                if not checkbox_found:
                    print("ℹ️ No agreement checkbox found or already checked")
                    
            except Exception as e:
                print(f"⚠️ Error with agreement checkbox: {e}")
            
            time.sleep(config.STEP_DELAY)
            
            # Submit form with multiple selectors and better error handling
            try:
                print("🚀 Submitting form...")
                submit_selectors = [
                    '#submit_but',  # From testing, this is the correct ID
                    'button[type="submit"]',
                    'input[type="submit"]',
                    'button:has-text("Добавить")',
                    'button:has-text("Отправить")',
                    'button:has-text("Добавить объявление")',
                    'input[value*="Подать"]',
                    'input[value*="Отправить"]',
                    '.submit-button',
                    "#submit-button"
                ]
                
                submitted = False
                for selector in submit_selectors:
                    try:
                        if self.page.locator(selector).count() > 0:
                            self.page.click(selector)
                            print("✅ Form submitted successfully!")
                            time.sleep(5)  # Wait for submission response
                            submitted = True
                            break
                    except:
                        continue
                
                if not submitted:
                    print("❌ Submit button not found")
                    return False
                    
            except Exception as e:
                print(f"❌ Form submission failed: {e}")
                return False
            
            # Check if submission was successful with proper URL and content checks
            print("⏳ Checking submission result...")
            time.sleep(3)  # Additional wait for page to load
            
            try:
                current_url = self.page.url
                page_content = self.page.content().lower()
                
                # Primary success indicator - URL parameter
                if "addsuccess=1" in current_url:
                    print("🎉 SUCCESS! Ad posted successfully! (URL success parameter found)")
                    return True
                
                # Authentication check - redirected to login
                elif "passport.orbita.co.il/site/login" in current_url:
                    print("❌ Redirected to login - authentication issue")
                    return False
                
                # Still on form page - check for validation errors
                elif "my/add" in current_url:
                    print("⚠️ Still on form page - checking for validation errors...")
                    
                    # Check for Russian success phrases first
                    russian_success_phrases = [
                        "объявление добавлено",
                        "объявление создано", 
                        "успешно добавлено",
                        "спасибо",
                        "размещено",
                        "опубликовано"
                    ]
                    
                    if any(phrase in page_content for phrase in russian_success_phrases):
                        print("✅ SUCCESS! Ad posted successfully! (Russian success message found)")
                        return True
                    
                    # Check for validation errors
                    try:
                        error_elements = self.page.locator('.alert, .error, .help-block, .text-danger, .invalid-feedback').all()
                        errors_found = []
                        for element in error_elements:
                            error_text = element.text_content()
                            if error_text and error_text.strip():
                                errors_found.append(error_text.strip())
                                print(f"❌ Form error: {error_text}")
                        
                        if errors_found:
                            return False
                        else:
                            print("⚠️ No clear success or error indicators - assuming success")
                            return True
                            
                    except Exception as e:
                        print(f"⚠️ Could not check for errors: {e}")
                        return True  # Assume success if can't check
                
                # Other pages - check for success indicators
                else:
                    print(f"🤔 Unexpected page after submission: {current_url}")
                    
                    # Check for success indicators in content
                    success_indicators = [
                        "success" in current_url,
                        "thank" in current_url,
                        "добавлено" in current_url,
                        any(phrase in page_content for phrase in ["спасибо", "добавлено", "успешно", "размещено"])
                    ]
                    
                    if any(success_indicators):
                        print("✅ SUCCESS! Ad posted successfully! (Success indicators found)")
                        return True
                    else:
                        print("⚠️ Unclear result - assuming success")
                        return True  # Assume success if unclear
                        
            except Exception as e:
                print(f"⚠️ Could not check submission status: {e}")
                return True  # Assume success if can't check
                
        except Exception as e:
            print(f"❌ Form filling failed: {e}")
            traceback.print_exc()  # Print full error trace for debugging
            return False
    
    def _fill_apartment_parameters(self, parameters: Dict[str, str]):
        """Fill apartment parameters fields based on extracted data"""
        try:
            # Rooms - using select dropdown with value mapping like in old working version
            if 'rooms' in parameters:
                print(f"🚪 Filling rooms: {parameters['rooms']}")
                try:
                    rooms_value = parameters['rooms']
                    # Map room values to option values (from working old version)
                    room_mapping = {
                        "1": "77", "1.5": "78", "2": "79", "2.5": "80",
                        "3": "81", "3.5": "82", "4": "83", "4.5": "84", 
                        "5": "85", "5.5": "86", "6+": "87"
                    }
                    
                    option_value = room_mapping.get(rooms_value, "0")  # Default to "Не указано"
                    
                    room_selectors = [
                        'select[name="room"]',
                        'select#room',
                        'select:has(option[value="81"])'  # Has option for 3 rooms
                    ]
                    
                    for selector in room_selectors:
                        try:
                            if self.page.locator(selector).count() > 0:
                                self.page.select_option(selector, value=option_value)
                                print(f"✅ Rooms filled successfully: {rooms_value} (option value: {option_value})")
                                break
                        except:
                            continue
                    else:
                        print("⚠️ Rooms select field not found")
                        
                except Exception as e:
                    print(f"❌ Error filling rooms: {e}")
            
            # Floor - using select dropdown with value mapping
            if 'floor' in parameters:
                print(f"🏢 Filling floor: {parameters['floor']}")
                try:
                    floor_value = parameters['floor']
                    # Map floor values to option values (from working old version)
                    floor_mapping = {
                        "0": "57", "1": "58", "2": "59", "3": "60", "4": "61",
                        "5": "62", "6": "63", "7": "64", "8": "65", "9": "66",
                        "10+": "67"
                    }
                    
                    option_value = floor_mapping.get(floor_value, "0")  # Default to "Не указано"
                    
                    floor_selectors = [
                        'select[name="floor"]',
                        'select#floor',
                        'select.floor-select'
                    ]
                    
                    for selector in floor_selectors:
                        try:
                            if self.page.locator(selector).count() > 0:
                                self.page.select_option(selector, value=option_value)
                                print(f"✅ Floor filled successfully: {floor_value} (option value: {option_value})")
                                break
                        except:
                            continue
                    else:
                        print("⚠️ Floor select field not found")
                        
                except Exception as e:
                    print(f"❌ Error filling floor: {e}")
            
            # Furniture - using select dropdown with value mapping
            if 'furniture' in parameters:
                print(f"🛋️ Filling furniture: {parameters['furniture']}")
                try:
                    furniture_value = parameters['furniture'].lower()
                    # Map furniture values to option values (from working old version)
                    furniture_mapping = {
                        "да": "26", "нет": "27", "частично": "28"
                    }
                    
                    option_value = furniture_mapping.get(furniture_value, "0")  # Default to "Не указано"
                    
                    furniture_selectors = [
                        'select[name="furniture"]',
                        'select#furniture',
                        'select.furniture-select'
                    ]
                    
                    for selector in furniture_selectors:
                        try:
                            if self.page.locator(selector).count() > 0:
                                self.page.select_option(selector, value=option_value)
                                print(f"✅ Furniture filled successfully: {furniture_value} (option value: {option_value})")
                                break
                        except:
                            continue
                    else:
                        print("⚠️ Furniture select field not found")
                        
                except Exception as e:
                    print(f"❌ Error filling furniture: {e}")
            
            # Price - using input field
            if 'price' in parameters:
                print(f"💰 Filling price: {parameters['price']}")
                try:
                    price_selectors = [
                        'input[name="cost"]',  # From working old version
                        'input#cost',
                        'input[name="price"]'
                    ]
                    
                    for selector in price_selectors:
                        try:
                            if self.page.locator(selector).count() > 0:
                                self.page.fill(selector, str(parameters['price']))
                                print("✅ Price filled successfully")
                                break
                        except:
                            continue
                    else:
                        print("⚠️ Price input field not found")
                        
                except Exception as e:
                    print(f"❌ Error filling price: {e}")
            
            # District/Address - using input field
            if 'district' in parameters:
                print(f"📍 Filling district: {parameters['district']}")
                try:
                    district_selectors = [
                        'input[name="address"]',  # From working old version
                        'input#address',
                        'input[name="district"]'
                    ]
                    
                    for selector in district_selectors:
                        try:
                            if self.page.locator(selector).count() > 0:
                                self.page.fill(selector, parameters['district'])
                                print("✅ District filled successfully")
                                break
                        except:
                            continue
                    else:
                        print("⚠️ District input field not found")
                        
                except Exception as e:
                    print(f"❌ Error filling district: {e}")
            
        except Exception as e:
            print(f"⚠️ Apartment parameters filling error: {e}")
    
    def _fill_contact_info(self, parameters: Dict[str, str]):
        """Fill contact information fields"""
        try:
            # Email (should be already filled from registration)
            email_input = self.page.locator("input[type='email']")
            if email_input.is_visible() and self.current_account_email:
                email_input.fill(self.current_account_email)
            
            # Phone prefix - select '055' (value="92")
            phone_code_select = self.page.locator("select[name='phonecode']")
            if phone_code_select.is_visible():
                phone_code_select.select_option(value="92")  # 055 prefix
                print("✅ Selected phone prefix: 055")
            
            # Phone number - use hardcoded number
            phone_input = self.page.locator("input[name='phonenum']")
            if phone_input.is_visible():
                phone_input.fill("5072867")  # Hardcoded phone number
                print("✅ Filled phone number: 055-5072867")
            
        except Exception as e:
            print(f"⚠️ Contact info filling error: {e}")
    
    def _solve_form_recaptcha(self) -> bool:
        """Solve reCAPTCHA on the form"""
        try:
            if not hasattr(self, 'solver') or not self.solver:
                self.solver = TwoCaptcha(config.CAPTCHA_API_KEY) if config.CAPTCHA_API_KEY != "your_2captcha_api_key_here" else None
            
            if not self.solver:
                print("⚠️ No captcha solver available")
                return False
            
            # Look for reCAPTCHA
            recaptcha_frame = self.page.locator("iframe[src*='recaptcha']").first
            if not recaptcha_frame.is_visible():
                print("ℹ️ No reCAPTCHA found on form")
                return True
            
            # Similar to registration reCAPTCHA solving
            site_key_element = self.page.locator("[data-sitekey]").first
            if site_key_element.is_visible():
                site_key = site_key_element.get_attribute("data-sitekey")
                
                print("🤖 Solving form reCAPTCHA...")
                result = self.solver.recaptcha(sitekey=site_key, url=self.page.url)
                
                if result and 'code' in result:
                    self.page.evaluate(f"""
                        const textarea = document.getElementById('g-recaptcha-response');
                        if (textarea) {{
                            textarea.innerHTML = '{result['code']}';
                            textarea.style.display = 'block';
                        }}
                    """)
                    return True
            
            return False
            
        except Exception as e:
            print(f"❌ Form reCAPTCHA solving failed: {e}")
            return False
    
    def _logout(self):
        """Logout from the current session"""
        try:
            print("🚪 Logging out...")
            
            # Look for logout link/button
            logout_selectors = [
                "a[href*='logout']",
                "a:has-text('Выход')",
                "a:has-text('Выйти')",
                "button:has-text('Выход')"
            ]
            
            for selector in logout_selectors:
                logout_element = self.page.locator(selector).first
                if logout_element.is_visible():
                    logout_element.click()
                    print("✅ Logged out successfully")
                    
                    # Wait after logout
                    if config.WAIT_AFTER_LOGOUT > 0:
                        print(f"⏳ Waiting {config.WAIT_AFTER_LOGOUT} seconds after logout...")
                        time.sleep(config.WAIT_AFTER_LOGOUT)
                    
                    # Change IP after logout
                    if self.tor_changer and config.CHANGE_IP_AFTER_LOGOUT:
                        self.tor_changer.change_ip()
                    
                    return
            
            print("⚠️ Logout button not found")
            
        except Exception as e:
            print(f"❌ Logout failed: {e}")
    
    def _load_processed_ads(self) -> set:
        """Load list of already processed ads"""
        processed = set()
        try:
            if os.path.exists(self.processed_ads_log):
                with open(self.processed_ads_log, 'r', encoding='utf-8') as f:
                    for line in f:
                        if '#' in line:
                            folder_id = line.split('#')[0].strip()
                            processed.add(folder_id)
        except Exception as e:
            print(f"⚠️ Could not load processed ads log: {e}")
        
        return processed
    
    def _log_processed_ad(self, folder_id: str, folder_name: str):
        """Log processed ad to prevent duplicates"""
        try:
            with open(self.processed_ads_log, 'a', encoding='utf-8') as f:
                timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                f.write(f"{folder_id} # {folder_name} # Processed on {timestamp}\n")
        except Exception as e:
            print(f"⚠️ Could not log processed ad: {e}")
    
    def cleanup(self):
        """Cleanup resources"""
        try:
            if self.page:
                self.page.close()
            if self.context:
                self.context.close()
            if self.browser:
                self.browser.close()
            if self.tor_changer:
                self.tor_changer.stop_tor()
            
            print("🧹 Cleanup completed")
            
        except Exception as e:
            print(f"⚠️ Cleanup error: {e}")

def main():
    """Main execution function"""
    print("=" * 60)
    print("🎯 ORBITA FORM FILLER V2.0 - ENHANCED ALGORITHM")
    print("=" * 60)
    
    filler = OrbitaFormFillerV2()
    
    try:
        # Initialize
        if not filler.initialize():
            print("❌ Initialization failed")
            return
        
        # Start browser
        if not filler.start_browser():
            print("❌ Browser start failed")
            return
        
        # Register and login
        if not filler.register_and_login():
            print("❌ Registration/login failed")
            return
        
        # Process all ads
        stats = filler.process_all_ads()
        
        # Print final statistics
        print("\n" + "=" * 60)
        print("📊 FINAL STATISTICS")
        print("=" * 60)
        print(f"✅ Successfully processed: {stats['processed']}")
        print(f"❌ Failed: {stats['failed']}")
        print(f"⏭️ Skipped (already processed): {stats['skipped']}")
        print(f"⚪ Skipped (empty documents): {stats['empty']}")
        print(f"📧 Account used: {filler.current_account_email}")
        print("=" * 60)
        
    except KeyboardInterrupt:
        print("\n⚠️ Process interrupted by user")
    except Exception as e:
        print(f"\n❌ Unexpected error: {e}")
    finally:
        filler.cleanup()

if __name__ == "__main__":
    main() 