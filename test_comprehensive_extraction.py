#!/usr/bin/env python3
"""
Comprehensive test for both Google Doc and .docx text extraction + OpenAI
"""

import os
import io
import tempfile
from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
import openai
import config

def setup_drive_api():
    """Setup Google Drive API authentication"""
    creds = None
    if os.path.exists('token.json'):
        creds = Credentials.from_authorized_user_file('token.json')
    
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
    
    service = build('drive', 'v3', credentials=creds)
    print("✅ Google Drive API connected successfully!")
    return service

def test_google_doc_extraction(service, file_id, filename):
    """Test Google Doc text extraction"""
    print(f"\n📄 Testing Google Doc extraction for: {filename}")
    print("=" * 50)
    
    try:
        request = service.files().export_media(
            fileId=file_id, 
            mimeType='text/plain'
        )
        
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, request)
        
        done = False
        while done is False:
            status, done = downloader.next_chunk()
        
        fh.seek(0)
        text = fh.read().decode('utf-8')
        
        print(f"✅ Extracted {len(text)} characters from Google Doc")
        if text:
            print(f"📝 First 300 characters: {text[:300]}...")
        
        return text
        
    except Exception as e:
        print(f"❌ Error extracting Google Doc: {e}")
        return ""

def test_docx_extraction(service, file_id, filename):
    """Test .docx text extraction with improved method"""
    print(f"\n📄 Testing .docx extraction for: {filename}")
    print("=" * 50)
    
    try:
        # Download the .docx file
        request = service.files().get_media(fileId=file_id)
        
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, request)
        
        done = False
        while done is False:
            status, done = downloader.next_chunk()
        
        print(f"✅ Downloaded {len(fh.getvalue())} bytes")
        
        # Extract text using python-docx
        fh.seek(0)
        
        try:
            from docx import Document
            doc = Document(fh)
            
            # Extract all text content
            text_content = []
            
            # From paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # From tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text.strip())
            
            # From headers/footers if available
            for section in doc.sections:
                if hasattr(section, 'header') and section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            text_content.append(paragraph.text.strip())
                
                if hasattr(section, 'footer') and section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            text_content.append(paragraph.text.strip())
            
            final_text = '\n'.join(text_content)
            print(f"✅ Extracted {len(final_text)} characters using python-docx")
            
            if final_text:
                print(f"📝 First 300 characters: {final_text[:300]}...")
                return final_text
            else:
                # Try alternative XML extraction
                print("🔍 Trying alternative XML extraction...")
                return extract_docx_via_xml(fh)
            
        except Exception as e:
            print(f"⚠️ python-docx method failed: {e}")
            # Try alternative XML extraction
            return extract_docx_via_xml(fh)
            
    except Exception as e:
        print(f"❌ Error downloading .docx file: {e}")
        return ""

def extract_docx_via_xml(file_handle):
    """Alternative .docx extraction using XML parsing"""
    try:
        file_handle.seek(0)
        import zipfile
        import re
        import xml.etree.ElementTree as ET
        
        with zipfile.ZipFile(file_handle, 'r') as zip_file:
            if 'word/document.xml' in zip_file.namelist():
                xml_content = zip_file.read('word/document.xml').decode('utf-8')
                
                # Parse XML properly
                try:
                    root = ET.fromstring(xml_content)
                    
                    # Find all text elements
                    text_elements = []
                    for elem in root.iter():
                        if elem.tag.endswith('}t') and elem.text:
                            text_elements.append(elem.text)
                    
                    if text_elements:
                        alt_text = ' '.join(text_elements)
                        print(f"✅ XML extraction: {len(alt_text)} characters")
                        print(f"📝 First 300 characters: {alt_text[:300]}...")
                        return alt_text
                        
                except ET.ParseError:
                    # Fallback to regex
                    text_matches = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', xml_content)
                    if text_matches:
                        alt_text = ' '.join(text_matches)
                        print(f"✅ Regex extraction: {len(alt_text)} characters")
                        print(f"📝 First 300 characters: {alt_text[:300]}...")
                        return alt_text
        
        print("❌ No text found in .docx file")
        return ""
        
    except Exception as e:
        print(f"❌ XML extraction failed: {e}")
        return ""

def test_openai_extraction(text, filename):
    """Test OpenAI parameter extraction"""
    print(f"\n🤖 Testing OpenAI extraction for: {filename}")
    print("=" * 50)
    
    if not text.strip():
        print("❌ No text to process")
        return {}
    
    try:
        # Initialize OpenAI client
        client = openai.OpenAI(api_key=config.OPENAI_API_KEY)
        
        prompt = """
Ты эксперт по анализу объявлений о продаже квартир в Израиле. 

Проанализируй следующий текст объявления и извлеки параметры в JSON формате:

1. "rooms" - количество комнат (например: "3", "4", "5", "2.5", "4.5")
2. "floor" - этаж (например: "2", "5", "10", "высокий этаж" -> "10+")  
3. "furniture" - мебель (если упоминается мебель/обставлена -> "да", иначе -> "нет")
4. "price" - цена в шекелях (только цифры, без валюты)
5. "district" - район Ришон-ле-Циона (ищи названия районов как "НАХЛАД ИУДА", "РЕМЕЗ", "НЕВЕ ДЕНЯ" и т.д.)

Если параметр не найден, не включай его в ответ.

Верни только JSON без дополнительного текста.

Текст объявления:
"""
        
        print("📤 Sending to OpenAI...")
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an expert at extracting apartment parameters from Russian real estate ads."},
                {"role": "user", "content": f"{prompt}\n\n{text[:2000]}"}  # Limit text length
            ],
            max_tokens=500,
            temperature=0.1
        )
        
        content = response.choices[0].message.content.strip()
        print(f"📥 OpenAI response: {content}")
        
        # Try to parse JSON
        import json
        try:
            # Clean up response if it has markdown formatting
            if content.startswith('```'):
                content = content.split('```')[1]
                if content.startswith('json'):
                    content = content[4:]
            
            parameters = json.loads(content)
            print(f"✅ Parsed parameters: {parameters}")
            return parameters
        except json.JSONDecodeError:
            print(f"⚠️ Could not parse as JSON, raw response: {content}")
            return {}
            
    except Exception as e:
        print(f"❌ OpenAI extraction failed: {e}")
        return {}

def find_and_test_documents(service):
    """Find and test both Google Docs and .docx files"""
    print("\n🔍 Finding documents to test...")
    
    # Find the ПРОДАЖА folder
    path = "Real estate/Ришон Лецион/ПРОДАЖА"
    path_parts = path.split('/')
    current_folder_id = 'root'
    
    for part in path_parts:
        query = f"name='{part}' and parents in '{current_folder_id}' and mimeType='application/vnd.google-apps.folder'"
        results = service.files().list(q=query).execute()
        folders = results.get('files', [])
        if folders:
            current_folder_id = folders[0]['id']
        else:
            print(f"❌ Could not find folder: {part}")
            return
    
    # Get all ad folders
    query = f"parents in '{current_folder_id}' and mimeType='application/vnd.google-apps.folder'"
    results = service.files().list(q=query).execute()
    folders = results.get('files', [])
    
    print(f"✅ Found {len(folders)} ad folders")
    
    # Test first few folders
    for i, folder in enumerate(folders[:3]):
        print(f"\n📁 Testing folder {i+1}: {folder['name']}")
        print("=" * 60)
        
        # Get folder contents
        query = f"parents in '{folder['id']}'"
        results = service.files().list(q=query).execute()
        files = results.get('files', [])
        
        # Find text documents
        google_docs = []
        docx_files = []
        
        for file in files:
            if file['mimeType'] == 'application/vnd.google-apps.document':
                google_docs.append(file)
            elif file['mimeType'] == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                docx_files.append(file)
        
        print(f"📊 Found: {len(google_docs)} Google Docs, {len(docx_files)} .docx files")
        
        # Test Google Docs
        for doc in google_docs:
            text = test_google_doc_extraction(service, doc['id'], doc['name'])
            if text:
                parameters = test_openai_extraction(text, doc['name'])
        
        # Test .docx files
        for doc in docx_files:
            text = test_docx_extraction(service, doc['id'], doc['name'])
            if text:
                parameters = test_openai_extraction(text, doc['name'])
        
        if not google_docs and not docx_files:
            print("❌ No text documents found in this folder")

def main():
    print("🧪 COMPREHENSIVE TEXT EXTRACTION TEST")
    print("📄 Testing both Google Docs and .docx files")
    print("=" * 60)
    
    # Setup API
    service = setup_drive_api()
    
    # Find and test documents
    find_and_test_documents(service)
    
    print("\n🏁 Comprehensive test completed!")

if __name__ == "__main__":
    main() 