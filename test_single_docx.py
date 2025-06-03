#!/usr/bin/env python3
"""
Test specific .docx file that was failing
"""

import os
import io
from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseDownload
import zipfile
import re
import xml.etree.ElementTree as ET

def setup_drive_api():
    """Setup Google Drive API authentication"""
    creds = None
    if os.path.exists('token.json'):
        creds = Credentials.from_authorized_user_file('token.json')
    
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
    
    service = build('drive', 'v3', credentials=creds)
    return service

def analyze_docx_structure(service, file_id, filename):
    """Deep analysis of .docx file structure"""
    print(f"\n🔍 Deep analysis of: {filename}")
    print("=" * 60)
    
    try:
        # Download the .docx file
        request = service.files().get_media(fileId=file_id)
        
        fh = io.BytesIO()
        downloader = MediaIoBaseDownload(fh, request)
        
        done = False
        while done is False:
            status, done = downloader.next_chunk()
        
        print(f"✅ Downloaded {len(fh.getvalue())} bytes")
        
        # Analyze ZIP structure
        fh.seek(0)
        with zipfile.ZipFile(fh, 'r') as zip_file:
            print(f"\n📦 ZIP structure:")
            for name in zip_file.namelist()[:15]:  # First 15 files
                info = zip_file.getinfo(name)
                print(f"   {name} ({info.file_size} bytes)")
            
            # Focus on document.xml
            if 'word/document.xml' in zip_file.namelist():
                xml_content = zip_file.read('word/document.xml').decode('utf-8', errors='ignore')
                print(f"\n📄 document.xml size: {len(xml_content)} characters")
                
                # Show first 1000 chars
                print(f"\n📝 XML preview (first 1000 chars):")
                print("-" * 50)
                print(xml_content[:1000])
                print("-" * 50)
                
                # Try different extraction methods
                print(f"\n🔧 Extraction attempts:")
                
                # Method 1: Simple regex
                text_matches = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', xml_content)
                if text_matches:
                    simple_text = ' '.join(text_matches)
                    print(f"✅ Simple regex: {len(simple_text)} chars")
                    print(f"   Preview: {simple_text[:200]}...")
                else:
                    print("❌ Simple regex: No matches")
                
                # Method 2: More comprehensive regex
                all_text = re.findall(r'<w:t[^>]*>([^<]*)</w:t>', xml_content)
                if all_text:
                    comprehensive_text = ' '.join(filter(None, all_text))
                    print(f"✅ Comprehensive regex: {len(comprehensive_text)} chars")
                    print(f"   Preview: {comprehensive_text[:200]}...")
                else:
                    print("❌ Comprehensive regex: No matches")
                
                # Method 3: XML parsing
                try:
                    root = ET.fromstring(xml_content)
                    text_elements = []
                    for elem in root.iter():
                        if elem.tag.endswith('}t') and elem.text:
                            text_elements.append(elem.text)
                    
                    if text_elements:
                        xml_text = ' '.join(text_elements)
                        print(f"✅ XML parsing: {len(xml_text)} chars")
                        print(f"   Preview: {xml_text[:200]}...")
                    else:
                        print("❌ XML parsing: No text elements")
                        
                except ET.ParseError as e:
                    print(f"❌ XML parsing failed: {e}")
                
                # Method 4: Check for specific namespaces
                namespaces = re.findall(r'xmlns:(\w+)="([^"]*)"', xml_content)
                print(f"\n📋 XML namespaces found:")
                for prefix, uri in namespaces[:5]:
                    print(f"   {prefix}: {uri}")
            
            # Check styles.xml for formatting clues
            if 'word/styles.xml' in zip_file.namelist():
                styles_content = zip_file.read('word/styles.xml').decode('utf-8', errors='ignore')
                print(f"\n🎨 styles.xml size: {len(styles_content)} characters")
        
        # Try python-docx analysis
        fh.seek(0)
        try:
            from docx import Document
            doc = Document(fh)
            
            print(f"\n📊 python-docx analysis:")
            print(f"   Paragraphs: {len(doc.paragraphs)}")
            print(f"   Tables: {len(doc.tables)}")
            print(f"   Sections: {len(doc.sections)}")
            
            # Check each paragraph
            for i, para in enumerate(doc.paragraphs[:10]):
                print(f"   Para {i+1}: '{para.text}' (style: {para.style.name if para.style else 'None'})")
                
                # Check runs within paragraph
                for j, run in enumerate(para.runs):
                    if run.text.strip():
                        print(f"     Run {j+1}: '{run.text}' (bold: {run.bold}, italic: {run.italic})")
            
        except Exception as e:
            print(f"❌ python-docx analysis failed: {e}")
            
    except Exception as e:
        print(f"❌ Error analyzing file: {e}")

def main():
    print("🔍 SPECIFIC DOCX FILE ANALYSIS")
    print("=" * 60)
    
    service = setup_drive_api()
    
    # Test the problematic file
    file_id = "1OBQr4XDqdGkNQ1xHn9VSrpNjBz-3LM9I"  # бэн аришон.docx
    filename = "бэн аришон.docx"
    
    analyze_docx_structure(service, file_id, filename)

if __name__ == "__main__":
    main() 